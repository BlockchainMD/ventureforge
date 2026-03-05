"""DOCX export via python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from ventureforge.core.config import get_settings
from ventureforge.utils.logger import get_logger

logger = get_logger()


def export_docx(markdown_content: str, filename: str, title: str = "VentureForge Report") -> Path:
    """Export Markdown content as a DOCX file."""
    output_dir = Path(get_settings().output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    filepath = output_dir / filename

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    doc.add_heading(title, level=0)

    for line in markdown_content.split("\n"):
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            # Skip markdown table rows
            continue
        elif line.startswith("---"):
            doc.add_page_break()
        else:
            # Handle inline bold
            text = line.replace("**", "")
            doc.add_paragraph(text)

    doc.save(str(filepath))
    logger.info("docx_exported", path=str(filepath))
    return filepath
